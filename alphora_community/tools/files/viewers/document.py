"""
Word 文档查看器 - 处理 .docx/.doc 文件
"""
import os
import re
from typing import Optional, List, Dict, Any, Tuple

from ..utils.common import get_file_info, truncate_text


class DocumentViewer:
    """Word 文档查看器"""
    
    SUPPORTED_EXTENSIONS = {'.docx', '.doc'}
    
    def __init__(self, file_path: str):
        """
        初始化查看器
        
        Args:
            file_path: 文件路径
        """
        self.file_path = file_path
        self.file_info = get_file_info(file_path)
        self.ext = self.file_info['extension']
        
    def view(
        self,
        purpose: str = "preview",
        keyword: Optional[str] = None,
        max_lines: int = 100,
        page_number: Optional[int] = None,
    ) -> str:
        """
        查看 Word 文档内容
        
        Args:
            purpose: 查看目的（preview/structure/search）
            keyword: 搜索关键词
            max_lines: 最大返回行数
            page_number: 页码（Word 不精确分页，仅用于大致定位）
            
        Returns:
            格式化的文档内容字符串
        """
        # 智能参数推断
        purpose, warnings = self._infer_and_validate_params(purpose, keyword)
        
        try:
            from docx import Document
        except ImportError:
            return "❌ 需要安装 python-docx 库：pip install python-docx"
        
        try:
            doc = Document(self.file_path)
        except Exception as e:
            return f"❌ 无法打开文档: {e}"
        
        # 提取内容
        paragraphs = self._extract_paragraphs(doc)
        tables = self._extract_tables(doc)
        
        if purpose == "structure":
            return self._get_structure(paragraphs, tables, warnings)
        elif purpose == "search":
            return self._search(paragraphs, tables, keyword, max_lines, warnings)
        else:  # preview
            return self._preview(paragraphs, tables, max_lines, warnings)
    
    def _infer_and_validate_params(
        self,
        purpose: str,
        keyword: Optional[str]
    ) -> Tuple[str, List[str]]:
        """智能推断和校验参数"""
        warnings = []
        
        if keyword and purpose != "search":
            warnings.append(f"⚠️ 检测到 keyword='{keyword}'，已自动切换为 search 模式")
            purpose = "search"
        
        if purpose == "search" and not keyword:
            warnings.append("⚠️ search 模式需要 keyword 参数，已切换为 preview 模式")
            purpose = "preview"
            
        return purpose, warnings
    
    def _extract_paragraphs(self, doc) -> List[Dict[str, Any]]:
        """提取所有段落"""
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else "Normal"
                paragraphs.append({
                    'text': text,
                    'style': style,
                    'is_heading': 'Heading' in style,
                    'heading_level': self._get_heading_level(style)
                })
        return paragraphs
    
    def _extract_tables(self, doc) -> List[List[List[str]]]:
        """提取所有表格"""
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables
    
    def _get_heading_level(self, style: str) -> int:
        """从样式名获取标题级别"""
        match = re.search(r'Heading\s*(\d)', style)
        return int(match.group(1)) if match else 0
    
    def _format_header(self, paragraphs: List, tables: List, warnings: List[str]) -> str:
        """格式化文件头信息"""
        lines = [
            f"📄 文件: {self.file_info['name']}",
            f"📦 大小: {self.file_info['size_human']}",
            f"📋 段落数: {len(paragraphs)} | 表格数: {len(tables)}",
        ]
        
        if warnings:
            lines.append("")
            for w in warnings:
                lines.append(w)
        
        return '\n'.join(lines)
    
    def _get_structure(
        self,
        paragraphs: List[Dict],
        tables: List,
        warnings: List[str]
    ) -> str:
        """获取文档结构"""
        lines = [
            self._format_header(paragraphs, tables, warnings),
            "",
            "【文档结构】"
        ]
        
        # 提取标题结构
        headings = [p for p in paragraphs if p['is_heading']]
        if headings:
            for h in headings[:30]:
                level = h['heading_level']
                indent = "  " * (level - 1) if level > 0 else ""
                text = truncate_text(h['text'], 60)
                lines.append(f"{indent}• {text}")
            
            if len(headings) > 30:
                lines.append(f"  ... 还有 {len(headings) - 30} 个标题")
        else:
            lines.append("  (文档没有标题结构)")
        
        # 表格概览
        if tables:
            lines.append("")
            lines.append("【表格列表】")
            for i, table in enumerate(tables[:10], 1):
                rows = len(table)
                cols = len(table[0]) if table else 0
                header = ' | '.join(table[0][:5]) if table and table[0] else '(空)'
                lines.append(f"  表格{i}: {rows}行 × {cols}列")
                lines.append(f"    表头: {truncate_text(header, 60)}")
            
            if len(tables) > 10:
                lines.append(f"  ... 还有 {len(tables) - 10} 个表格")
        
        return '\n'.join(lines)
    
    def _preview(
        self,
        paragraphs: List[Dict],
        tables: List,
        max_lines: int,
        warnings: List[str]
    ) -> str:
        """预览文档内容"""
        lines = [
            self._format_header(paragraphs, tables, warnings),
            "",
            "【内容预览】"
        ]
        
        count = 0
        for p in paragraphs:
            if count >= max_lines:
                break
            
            prefix = ""
            if p['is_heading']:
                level = p['heading_level']
                prefix = "#" * level + " " if level > 0 else "## "
            
            lines.append(f"{prefix}{p['text']}")
            count += 1
        
        if len(paragraphs) > max_lines:
            lines.append(f"\n... 还有 {len(paragraphs) - max_lines} 个段落未显示")
        
        # 显示表格摘要
        if tables:
            lines.append("")
            lines.append(f"📊 文档包含 {len(tables)} 个表格（使用 purpose='structure' 查看详情）")
        
        return '\n'.join(lines)
    
    def _search(
        self,
        paragraphs: List[Dict],
        tables: List,
        keyword: str,
        max_lines: int,
        warnings: List[str]
    ) -> str:
        """搜索文档内容"""
        results = []
        keyword_lower = keyword.lower()
        
        # 搜索段落
        for i, p in enumerate(paragraphs, 1):
            if keyword_lower in p['text'].lower():
                # 高亮关键词上下文
                text = p['text']
                idx = text.lower().find(keyword_lower)
                start = max(0, idx - 30)
                end = min(len(text), idx + len(keyword) + 30)
                context = text[start:end]
                if start > 0:
                    context = "..." + context
                if end < len(text):
                    context = context + "..."
                    
                results.append({
                    'type': 'paragraph',
                    'location': f"段落{i}",
                    'content': context
                })
        
        # 搜索表格
        for t_idx, table in enumerate(tables, 1):
            for r_idx, row in enumerate(table):
                for c_idx, cell in enumerate(row):
                    if keyword_lower in cell.lower():
                        results.append({
                            'type': 'table',
                            'location': f"表格{t_idx} 第{r_idx+1}行",
                            'content': truncate_text(cell, 100)
                        })
        
        # 格式化输出
        lines = [
            self._format_header(paragraphs, tables, warnings),
            "",
            f"🔍 搜索关键词: '{keyword}'",
            f"📋 找到 {len(results)} 处匹配"
        ]
        
        if not results:
            lines.append("")
            lines.append(f"未找到包含 '{keyword}' 的内容")
            return '\n'.join(lines)
        
        lines.append("")
        
        for i, r in enumerate(results[:max_lines], 1):
            lines.append(f"[{r['location']}] {r['content']}")
        
        if len(results) > max_lines:
            lines.append(f"\n... 还有 {len(results) - max_lines} 处匹配未显示")
        
        return '\n'.join(lines)
